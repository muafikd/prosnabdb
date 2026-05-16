"""
Services module for business logic, including cost calculation.
"""
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
from django.utils import timezone
from django.conf import settings
from django.db import transaction
from .models import (
    Equipment, EquipmentPhoto, PurchasePrice, Logistics, AdditionalPrices, ExchangeRate,
    CostCalculation, CommercialProposal, SystemSettings
)
import requests
import re
import os
from io import BytesIO
from django.core.cache import cache
from django.core.files.base import ContentFile
import logging
from django.utils.html import escape
from docx import Document
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docxtpl import DocxTemplate, InlineImage
from html.parser import HTMLParser

try:
    from PIL import Image
except ImportError:
    Image = None

logger = logging.getLogger(__name__)

class DocxHTMLParser(HTMLParser):
    def __init__(self, subdoc):
        super().__init__()
        self.subdoc = subdoc
        self.current_paragraph = self.subdoc.add_paragraph()
        self.current_paragraph.style.font.size = Pt(10)
        
        self.bold = False
        self.italic = False
        self.underline = False
        
    def handle_starttag(self, tag, attrs):
        if tag in ('b', 'strong'):
            self.bold = True
        elif tag in ('i', 'em'):
            self.italic = True
        elif tag == 'u':
            self.underline = True
        elif tag in ('p', 'div'):
            if len(self.current_paragraph.runs) > 0:
                self.current_paragraph = self.subdoc.add_paragraph()
                self.current_paragraph.style.font.size = Pt(10)
        elif tag == 'br':
            self.current_paragraph.add_run('\n')
            
    def handle_endtag(self, tag):
        if tag in ('b', 'strong'):
            self.bold = False
        elif tag in ('i', 'em'):
            self.italic = False
        elif tag == 'u':
            self.underline = False
        elif tag in ('p', 'div'):
            self.current_paragraph = self.subdoc.add_paragraph()
            self.current_paragraph.style.font.size = Pt(10)
            
    def handle_data(self, data):
        if data:
            run = self.current_paragraph.add_run(data)
            run.font.size = Pt(10)
            if self.bold: run.font.bold = True
            if self.italic: run.font.italic = True
            if self.underline: run.underline = True

class LinkConverterService:
    """Service for converting cloud storage links to direct download links."""
    
    @staticmethod
    def get_direct_link(url):
        """
        Convert cloud storage link to direct download link.
        Supported services: Google Drive, Yandex Disk.
        """
        if not url:
            return url
            
        if 'drive.google.com' in url:
            return LinkConverterService._convert_google_drive(url)
        elif 'disk.yandex' in url or 'yadi.sk' in url:
            return LinkConverterService._convert_yandex_disk(url)
            
        return url


class LinkConverterService:
    """Service for converting cloud storage links to direct download links."""
    
    @staticmethod
    def get_direct_link(url):
        """
        Convert cloud storage link to direct download link.
        Supported services: Google Drive, Yandex Disk.
        """
        if not url:
            return url
            
        if 'drive.google.com' in url:
            return LinkConverterService._convert_google_drive(url)
        elif 'disk.yandex' in url or 'yadi.sk' in url:
            return LinkConverterService._convert_yandex_disk(url)
            
        return url

    @staticmethod
    def _convert_google_drive(url):
        """Convert Google Drive link to direct download link."""
        try:
            # Extract file ID
            # Pattern 1: .../file/d/FILE_ID/...
            match = re.search(r'/file/d/([a-zA-Z0-9_-]+)', url)
            if match:
                file_id = match.group(1)
                return f"https://drive.google.com/uc?export=view&id={file_id}"
            
            # Pattern 2: ...?id=FILE_ID...
            match = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
            if match:
                file_id = match.group(1)
                return f"https://drive.google.com/uc?export=view&id={file_id}"
                
        except Exception as e:
            logger.warning(f"Failed to convert Google Drive link {url}: {e}")
            
        return url

    @staticmethod
    def _convert_yandex_disk(url):
        """Convert Yandex Disk link to direct download link with caching."""
        cache_key = f"yandex_disk_link_{url}"
        cached_link = cache.get(cache_key)
        
        if cached_link:
            return cached_link
            
        try:
            api_url = "https://cloud-api.yandex.net/v1/disk/public/resources/download"
            params = {'public_key': url}
            
            response = requests.get(api_url, params=params, timeout=5)
            
            if response.status_code == 200:
                data = response.json()
                href = data.get('href')
                if href:
                    # Cache for 30 minutes (1800 seconds)
                    cache.set(cache_key, href, timeout=1800)
                    return href
            else:
                logger.warning(f"Yandex API returned {response.status_code} for {url}")
                
        except Exception as e:
            logger.warning(f"Failed to convert Yandex Disk link {url}: {e}")
            
        return url


# Max dimension for optimized equipment photos (saves SSD space)
EQUIPMENT_PHOTO_MAX_PX = 1600
EQUIPMENT_PHOTO_JPEG_QUALITY = 85


class CloudImageImportService:
    """
    Service for downloading images from Google Drive / Yandex Disk,
    optimizing them (resize, JPEG) and storing locally in EquipmentPhoto.
    """
    SOURCE_GOOGLE = 'google'
    SOURCE_YANDEX = 'yandex'

    @staticmethod
    def detect_source(url):
        """Determine cloud source from URL. Returns SOURCE_GOOGLE, SOURCE_YANDEX, or None."""
        if not url or not isinstance(url, str):
            return None
        url = url.strip()
        if 'drive.google.com' in url:
            return CloudImageImportService.SOURCE_GOOGLE
        if 'disk.yandex' in url or 'yadi.sk' in url:
            return CloudImageImportService.SOURCE_YANDEX
        return None

    @staticmethod
    def _extract_google_file_id(url):
        """Extract Google Drive file ID from URL."""
        match = re.search(r'/file/d/([a-zA-Z0-9_-]+)', url)
        if match:
            return match.group(1)
        match = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
        if match:
            return match.group(1)
        return None

    @staticmethod
    def get_download_url(url):
        """
        Get direct download URL for the given cloud link.
        Google: /uc?export=download&id=FILE_ID
        Yandex: API public/resources/download -> href
        """
        if not url:
            return None
        url = url.strip()
        if CloudImageImportService.detect_source(url) == CloudImageImportService.SOURCE_GOOGLE:
            file_id = CloudImageImportService._extract_google_file_id(url)
            if file_id:
                return f"https://drive.google.com/uc?export=download&id={file_id}"
            return None
        if CloudImageImportService.detect_source(url) == CloudImageImportService.SOURCE_YANDEX:
            try:
                api_url = "https://cloud-api.yandex.net/v1/disk/public/resources/download"
                resp = requests.get(api_url, params={'public_key': url}, timeout=10)
                if resp.status_code == 200:
                    data = resp.json()
                    href = data.get('href')
                    if href:
                        return href
            except Exception as e:
                logger.warning("Yandex download URL failed for %s: %s", url[:80], e)
            return None
        return url

    @staticmethod
    def download_bytes(url, timeout=30):
        """Download image bytes from URL. Returns bytes or None."""
        if not url:
            return None
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (compatible; ProsnabDB/1.0)',
            }
            resp = requests.get(url, headers=headers, timeout=timeout, stream=True)
            resp.raise_for_status()
            # Google Drive may return HTML confirmation for large files; try to get direct link
            if 'drive.google.com' in url and 'text/html' in resp.headers.get('Content-Type', ''):
                # Try to extract confirm token and re-request
                match = re.search(r'confirm=([0-9A-Za-z_-]+)', resp.text)
                if match:
                    confirm = match.group(1)
                    download_url = url if '?' in url else f"{url}&confirm={confirm}"
                    resp = requests.get(download_url, headers=headers, timeout=timeout)
                    resp.raise_for_status()
            return resp.content
        except Exception as e:
            logger.warning("Download failed for %s: %s", url[:80], e)
            return None

    @staticmethod
    def optimize_image(raw_bytes):
        """
        Resize image so longest side is EQUIPMENT_PHOTO_MAX_PX and save as JPEG quality 85.
        Returns JPEG bytes or None.
        """
        if not Image or not raw_bytes:
            return raw_bytes
        try:
            img = Image.open(BytesIO(raw_bytes))
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            w, h = img.size
            if w > EQUIPMENT_PHOTO_MAX_PX or h > EQUIPMENT_PHOTO_MAX_PX:
                if w >= h:
                    new_w = EQUIPMENT_PHOTO_MAX_PX
                    new_h = int(h * EQUIPMENT_PHOTO_MAX_PX / w)
                else:
                    new_h = EQUIPMENT_PHOTO_MAX_PX
                    new_w = int(w * EQUIPMENT_PHOTO_MAX_PX / h)
                img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            out = BytesIO()
            img.save(out, format='JPEG', quality=EQUIPMENT_PHOTO_JPEG_QUALITY, optimize=True)
            return out.getvalue()
        except Exception as e:
            logger.warning("Optimize image failed: %s", e)
            return raw_bytes

    @staticmethod
    def import_and_save(equipment, url, name='', sort_order=0):
        """
        Download image from URL (Google/Yandex), optimize and save as EquipmentPhoto.
        Returns EquipmentPhoto instance or None on failure.
        """
        if not equipment or not url:
            return None
        url = url.strip()
        download_url = CloudImageImportService.get_download_url(url)
        if not download_url:
            # Not a supported cloud link; could be already local /media/...
            return None
        raw = CloudImageImportService.download_bytes(download_url)
        if not raw:
            return None
        jpeg_bytes = CloudImageImportService.optimize_image(raw)
        if not jpeg_bytes:
            return None
        try:
            photo = EquipmentPhoto(equipment=equipment, name=name or '', sort_order=sort_order)
            photo.image.save(
                f"{__import__('uuid').uuid4().hex}.jpg",
                ContentFile(jpeg_bytes),
                save=True
            )
            return photo
        except Exception as e:
            logger.warning("Save EquipmentPhoto failed for equipment %s: %s", equipment.equipment_id, e)
            return None


class CostCalculationService:
    """Service for calculating equipment cost."""
    
    @staticmethod
    def convert_currency(amount, from_currency, to_currency, date=None, proposal=None, manual_rate_value=None):
        """
        Конвертировать сумму из одной валюты в другую.
        
        Args:
            amount: Сумма для конвертации
            from_currency: Валюта "из"
            to_currency: Валюта "в"
            date: Дата курса (если None, используется сегодня)
            proposal: КП для получения корректировки курса
            manual_rate_value: Ручное значение курса (если указано, используется вместо поиска в БД)
        
        Returns:
            Decimal: Конвертированная сумма
        """
        if from_currency == to_currency:
            return Decimal(str(amount))
        
        # Использовать ручной курс, если указан
        if manual_rate_value is not None:
            rate_value = Decimal(str(manual_rate_value))
        else:
            # Получить курс валюты из БД
            rate = ExchangeRate.get_latest_rate(
                currency_from=from_currency,
                currency_to=to_currency,
                date=date,
                proposal=proposal
            )
            
            if not rate:
                raise ValueError(
                    f'Exchange rate not found for {from_currency}/{to_currency}'
                    + (f' on {date}' if date else '')
                    + '. Please provide exchange rate in the proposal or create it manually.'
                )
            
            rate_value = rate.rate_value
        
        # Конвертировать
        converted_amount = Decimal(str(amount)) * rate_value
        return converted_amount.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    
    @staticmethod
    def calculate_equipment_cost(
        equipment,
        purchase_price_id=None,
        logistics_id=None,
        additional_prices_id=None,
        exchange_rate_date=None,
        proposal=None,
        target_currency='KZT',
        manual_overrides=None
    ):
        """
        Рассчитать себестоимость оборудования по формуле:
        Себестоимость в KZT = (закупочная цена + логистика + склад + производственные расходы + доп.расходы)
        где все компоненты конвертируются в KZT через курс валюты из КП
        
        Args:
            equipment: Equipment объект
            purchase_price_id: ID конкретной цены закупки (если None, берется активная)
            logistics_id: ID конкретной логистики (если None, берется активная)
            additional_prices_id: ID дополнительных расходов (если None, не учитываются)
            exchange_rate_date: Дата курса валюты (если None, используется сегодня)
            proposal: CommercialProposal для получения корректировок курса
            target_currency: Целевая валюта для результата (по умолчанию KZT)
            manual_overrides: Словарь с ручными переопределениями значений
        
        Returns:
            dict: Словарь с результатами расчета
        """
        if exchange_rate_date is None:
            exchange_rate_date = timezone.now().date()
        
        manual_overrides = manual_overrides or {}
        
        # 1. Закупочная цена
        purchase_price_value = Decimal('0')
        purchase_price_currency = None
        purchase_price_source = None
        purchase_price_obj = None
        
        if purchase_price_id:
            try:
                purchase_price_obj = PurchasePrice.objects.get(
                    price_id=purchase_price_id,
                    equipment=equipment,
                    is_active=True
                )
            except PurchasePrice.DoesNotExist:
                raise ValueError(f'Purchase price with ID {purchase_price_id} not found')
        else:
            # Берем активную цену закупки
            purchase_price_obj = PurchasePrice.objects.filter(
                equipment=equipment,
                is_active=True
            ).order_by('-created_at').first()
        
        if purchase_price_obj:
            purchase_price_value = Decimal(str(purchase_price_obj.price))
            purchase_price_currency = purchase_price_obj.currency
            purchase_price_source = purchase_price_obj.source_type
        
        # Ручное переопределение
        if 'purchase_price' in manual_overrides:
            purchase_price_value = Decimal(str(manual_overrides['purchase_price']))
            if 'purchase_price_currency' in manual_overrides:
                purchase_price_currency = manual_overrides['purchase_price_currency']
        
        # 2. Логистика - суммируем все активные логистики (кроме склада kz_warehouse)
        logistics_cost = Decimal('0')
        logistics_currency = None
        logistics_route = None
        logistics_obj = None
        
        if logistics_id:
            try:
                logistics_obj = Logistics.objects.get(
                    logistics_id=logistics_id,
                    equipment=equipment,
                    is_active=True
                )
                logistics_cost = Decimal(str(logistics_obj.cost))
                logistics_currency = logistics_obj.currency
                logistics_route = logistics_obj.route_type
            except Logistics.DoesNotExist:
                raise ValueError(f'Logistics with ID {logistics_id} not found')
        else:
            # Суммируем все активные логистики (кроме склада kz_warehouse)
            active_logistics = Logistics.objects.filter(
                equipment=equipment,
                is_active=True
            ).exclude(route_type='kz_warehouse').order_by('-created_at')
            
            if active_logistics.exists():
                # Берем первую для route_type (для обратной совместимости)
                logistics_obj = active_logistics.first()
                logistics_route = logistics_obj.route_type
                logistics_currency = logistics_obj.currency
                # Суммируем логистики в основной валюте (конвертация всех в KZT будет позже)
                logistics_cost = sum(Decimal(str(l.cost)) for l in active_logistics if l.currency == logistics_currency)
        
        # Ручное переопределение
        if 'logistics_cost' in manual_overrides:
            logistics_cost = Decimal(str(manual_overrides['logistics_cost']))
            if 'logistics_currency' in manual_overrides:
                logistics_currency = manual_overrides['logistics_currency']
        
        # 3. Склад (из логистики с маршрутом kz_warehouse)
        warehouse_cost = Decimal('0')
        warehouse_currency = None
        
        warehouse_logistics = Logistics.objects.filter(
            equipment=equipment,
            route_type='kz_warehouse',
            is_active=True
        ).first()
        
        if warehouse_logistics:
            warehouse_cost = Decimal(str(warehouse_logistics.cost))
            warehouse_currency = warehouse_logistics.currency
        
        # Ручное переопределение
        if 'warehouse_cost' in manual_overrides:
            warehouse_cost = Decimal(str(manual_overrides['warehouse_cost']))
            if 'warehouse_currency' in manual_overrides:
                warehouse_currency = manual_overrides['warehouse_currency']
        
        # 4. Производственные расходы
        production_cost = Decimal('0')
        production_currency = None
        
        if equipment.equipment_manufacture_price:
            production_cost = Decimal(str(equipment.equipment_manufacture_price))
            production_currency = equipment.equipment_price_currency_type or 'KZT'
        
        # Ручное переопределение
        if 'production_cost' in manual_overrides:
            production_cost = Decimal(str(manual_overrides['production_cost']))
            if 'production_currency' in manual_overrides:
                production_currency = manual_overrides['production_currency']
        
        # 5. Дополнительные расходы (AdditionalPrices) - будут рассчитаны после базовой себестоимости
        # Сначала собираем информацию о доп.расходах
        additional_prices_list = []
        
        # Если есть proposal, проверяем EquipmentList для этого оборудования
        if proposal:
            from .models import EquipmentList, EquipmentListItem
            # Ищем EquipmentList в этом КП, который содержит это оборудование
            equipment_lists = EquipmentList.objects.filter(
                proposal=proposal
            ).prefetch_related('equipment_items_relation', 'additional_prices')
            
            for eq_list in equipment_lists:
                # Проверяем, есть ли это оборудование в списке
                if eq_list.equipment_items_relation.filter(equipment=equipment).exists():
                    # Собираем все дополнительные расходы
                    if eq_list.additional_prices.exists():
                        additional_prices_list.extend(eq_list.additional_prices.all())
                    # Поддержка старого поля для обратной совместимости
                    elif eq_list.additional_price:
                        additional_prices_list.append(eq_list.additional_price)
                    break
        
        # Если указан конкретный additional_prices_id
        if additional_prices_id:
            try:
                additional_prices_obj = AdditionalPrices.objects.get(price_id=additional_prices_id)
                additional_prices_list.append(additional_prices_obj)
            except AdditionalPrices.DoesNotExist:
                pass
        
        # Ручное переопределение - если указано, используем его вместо расчетных
        manual_additional_costs = None
        if 'additional_costs' in manual_overrides:
            manual_additional_costs = Decimal(str(manual_overrides['additional_costs']))
        
        # 6. Конвертация всех сумм в базовую валюту (KZT) через курс валюты
        service = CostCalculationService()
        
        # Получаем ручной курс валюты, если указан
        manual_rate_value = manual_overrides.get('exchange_rate_value')
        # Если курс не указан вручную, но есть proposal с курсом, используем его
        if not manual_rate_value and proposal and proposal.exchange_rate:
            manual_rate_value = proposal.exchange_rate
        
        # Определяем базовую валюту для расчета (валюта закупочной цены или KZT)
        base_currency = purchase_price_currency or 'KZT'
        
        # Конвертируем все компоненты в базовую валюту (KZT)
        # 1. Закупочная цена
        purchase_price_base = purchase_price_value
        if purchase_price_currency and purchase_price_currency != 'KZT':
            try:
                purchase_price_base = service.convert_currency(
                    purchase_price_value,
                    purchase_price_currency,
                    'KZT',
                    date=exchange_rate_date,
                    proposal=proposal,
                    manual_rate_value=manual_rate_value
                )
            except ValueError as e:
                # Если не удалось конвертировать, используем исходное значение
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Failed to convert purchase price from {purchase_price_currency} to KZT: {e}')
                purchase_price_base = purchase_price_value
        
        # 2. Логистика - суммируем все активные логистики, конвертируя каждую в KZT
        logistics_base = Decimal('0')
        if not logistics_id:
            # Берем все активные логистики (кроме склада)
            active_logistics = Logistics.objects.filter(
                equipment=equipment,
                is_active=True
            ).exclude(route_type='kz_warehouse')
            
            for log in active_logistics:
                log_cost = Decimal(str(log.cost))
                if log.currency == 'KZT':
                    logistics_base += log_cost
                else:
                    try:
                        # Конвертируем в KZT через курс из КП
                        log_cost_kzt = service.convert_currency(
                            log_cost,
                            log.currency,
                            'KZT',
                            date=exchange_rate_date,
                            proposal=proposal,
                            manual_rate_value=manual_rate_value
                        )
                        logistics_base += log_cost_kzt
                    except ValueError as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f'Failed to convert logistics from {log.currency} to KZT: {e}')
                        # Если не удалось конвертировать, пропускаем эту логистику
        elif logistics_cost > 0:
            # Если указана конкретная логистика или есть ручное переопределение
            if logistics_currency and logistics_currency != 'KZT':
                try:
                    logistics_base = service.convert_currency(
                        logistics_cost,
                        logistics_currency,
                        'KZT',
                        date=exchange_rate_date,
                        proposal=proposal,
                        manual_rate_value=manual_rate_value
                    )
                except ValueError as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f'Failed to convert logistics from {logistics_currency} to KZT: {e}')
                    logistics_base = logistics_cost
            else:
                logistics_base = logistics_cost
        
        # 3. Склад
        warehouse_base = warehouse_cost
        if warehouse_currency and warehouse_currency != 'KZT':
            try:
                warehouse_base = service.convert_currency(
                    warehouse_cost,
                    warehouse_currency,
                    'KZT',
                    date=exchange_rate_date,
                    proposal=proposal,
                    manual_rate_value=manual_rate_value
                )
            except ValueError as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Failed to convert warehouse cost from {warehouse_currency} to KZT: {e}')
                warehouse_base = warehouse_cost
        
        # 4. Производственные расходы
        production_base = production_cost
        if production_currency and production_currency != 'KZT':
            try:
                production_base = service.convert_currency(
                    production_cost,
                    production_currency,
                    'KZT',
                    date=exchange_rate_date,
                    proposal=proposal,
                    manual_rate_value=manual_rate_value
                )
            except ValueError as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Failed to convert production cost from {production_currency} to KZT: {e}')
                production_base = production_cost
        
        # 6. Рассчитываем базовую себестоимость (без доп.расходов)
        base_cost_kzt = (
            purchase_price_base +
            logistics_base +
            warehouse_base +
            production_base
        )
        
        # 7. Рассчитываем дополнительные расходы на основе базовой себестоимости
        additional_base = Decimal('0')
        
        # Если есть ручное переопределение, используем его
        if manual_additional_costs is not None:
            additional_base = manual_additional_costs
        else:
            # Рассчитываем дополнительные расходы по их типам
            for additional_price in additional_prices_list:
                value = Decimal(str(additional_price.price_parameter_value))
                
                if additional_price.value_type == 'percentage':
                    # Процент от базовой себестоимости
                    additional_base += base_cost_kzt * (value / Decimal('100'))
                elif additional_price.value_type == 'fixed':
                    # Фиксированная сумма
                    additional_base += value
                elif additional_price.value_type == 'coefficient':
                    # Коэффициент умножения базовой себестоимости
                    additional_base += base_cost_kzt * value
        
        # 8. Итоговая себестоимость в KZT (сумма всех компонентов)
        # Формула: Себестоимость = (закупочная цена + логистика + склад + производство) + доп.расходы
        total_cost_kzt = base_cost_kzt + additional_base
        
        # 8. Конвертация в целевую валюту (если не KZT)
        # Если целевая валюта не KZT, конвертируем через курс
        total_cost_target = total_cost_kzt
        if target_currency != 'KZT':
            try:
                # Для конвертации из KZT в целевую валюту
                # Если есть manual_rate_value из proposal, это курс из целевой валюты в KZT
                # Например, если proposal.currency_ticket = RUB, manual_rate_value = курс RUB/KZT
                # То для конвертации KZT -> RUB нужен обратный курс: 1 / manual_rate_value
                reverse_rate_value = None
                if manual_rate_value and proposal and proposal.currency_ticket == target_currency:
                    # Если целевая валюта совпадает с валютой КП, используем обратный курс
                    reverse_rate_value = Decimal('1') / Decimal(str(manual_rate_value))
                
                total_cost_target = service.convert_currency(
                    total_cost_kzt,
                    'KZT',
                    target_currency,
                    date=exchange_rate_date,
                    proposal=proposal,
                    manual_rate_value=reverse_rate_value
                )
            except ValueError as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Failed to convert total cost from KZT to {target_currency}: {e}')
                # Если не удалось конвертировать, оставляем в KZT
                total_cost_target = total_cost_kzt
        
        # 9. Получить курс валюты для сохранения
        exchange_rate = ExchangeRate.get_latest_rate(
            currency_from=purchase_price_currency or 'KZT',
            currency_to='KZT',
            date=exchange_rate_date,
            proposal=proposal
        )
        
        # 10. Формируем результат
        result = {
            'equipment_id': equipment.equipment_id,
            'equipment_name': equipment.equipment_name,
            'purchase_price': {
                'id': purchase_price_obj.price_id if purchase_price_obj else None,
                'value': float(purchase_price_value),
                'currency': purchase_price_currency,
                'source': purchase_price_source,
                'value_kzt': float(purchase_price_base),
            },
            'logistics': {
                'id': logistics_obj.logistics_id if logistics_obj else None,
                'cost': float(logistics_cost),
                'currency': logistics_currency,
                'route': logistics_route,
                'cost_kzt': float(logistics_base),
            },
            'warehouse': {
                'cost': float(warehouse_cost),
                'currency': warehouse_currency,
                'cost_kzt': float(warehouse_base),
            },
            'production': {
                'cost': float(production_cost),
                'currency': production_currency,
                'cost_kzt': float(production_base),
            },
            'additional_costs': {
                'id': additional_prices_id,
                'cost': float(additional_base),  # Используем рассчитанное значение в KZT
                'currency': 'KZT',  # Дополнительные расходы всегда в KZT после расчета
                'cost_kzt': float(additional_base),
            },
            'exchange_rate': {
                'id': exchange_rate.rate_id if exchange_rate else None,
                'value': float(exchange_rate.rate_value) if exchange_rate else 1.0,
                'from_currency': exchange_rate.currency_from if exchange_rate else 'KZT',
                'to_currency': exchange_rate.currency_to if exchange_rate else 'KZT',
                'date': str(exchange_rate.rate_date) if exchange_rate else str(exchange_rate_date),
            },
            'total_cost_kzt': float(total_cost_kzt),
            'total_cost_target_currency': float(total_cost_target),
            'target_currency': target_currency,
            'calculation_date': str(exchange_rate_date),
        }
        
        return result
    
    @staticmethod
    @transaction.atomic
    def save_calculation(
        equipment,
        calculation_result,
        proposal=None,
        created_by=None,
        is_manual_adjustment=False,
        notes=None
    ):
        """
        Сохранить расчет себестоимости в базу данных.
        
        Args:
            equipment: Equipment объект
            calculation_result: Результат расчета из calculate_equipment_cost
            proposal: CommercialProposal (опционально)
            created_by: User (опционально)
            is_manual_adjustment: Флаг ручной корректировки
            notes: Примечания
        
        Returns:
            CostCalculation: Созданный объект расчета
        """
        # Определить версию расчета
        if proposal:
            last_version = CostCalculation.objects.filter(
                equipment=equipment,
                proposal=proposal
            ).order_by('-calculation_version').first()
        else:
            last_version = CostCalculation.objects.filter(
                equipment=equipment,
                proposal__isnull=True
            ).order_by('-calculation_version').first()
        
        next_version = (last_version.calculation_version + 1) if last_version else 1
        
        # Создать запись расчета
        calculation = CostCalculation.objects.create(
            equipment=equipment,
            proposal=proposal,
            calculation_version=next_version,
            purchase_price_id=calculation_result['purchase_price']['id'],
            purchase_price_value=calculation_result['purchase_price']['value'],
            purchase_price_currency=calculation_result['purchase_price']['currency'],
            purchase_price_source=calculation_result['purchase_price']['source'],
            logistics_id=calculation_result['logistics']['id'],
            logistics_cost=calculation_result['logistics']['cost'],
            logistics_currency=calculation_result['logistics']['currency'],
            logistics_route=calculation_result['logistics']['route'],
            warehouse_cost=calculation_result['warehouse']['cost'],
            warehouse_currency=calculation_result['warehouse']['currency'],
            production_cost=calculation_result['production']['cost'],
            production_currency=calculation_result['production']['currency'],
            additional_costs=calculation_result['additional_costs']['cost'],
            additional_costs_currency=calculation_result['additional_costs']['currency'] or 'KZT',
            exchange_rate_id=calculation_result['exchange_rate']['id'],
            exchange_rate_value=calculation_result['exchange_rate']['value'],
            exchange_rate_from=calculation_result['exchange_rate']['from_currency'],
            exchange_rate_to=calculation_result['exchange_rate']['to_currency'],
            exchange_rate_date=calculation_result['exchange_rate']['date'],
            total_cost_base_currency=calculation_result['total_cost_kzt'],
            total_cost_kzt=calculation_result['total_cost_kzt'],
            calculation_details=calculation_result,
            is_manual_adjustment=is_manual_adjustment,
            notes=notes,
            created_by=created_by,
        )
        
        return calculation
    
    @staticmethod
    def calculate_proposal_total_price(cost_price, margin_percentage=None):
        """
        Рассчитать итоговую цену КП на основе себестоимости и маржи.
        
        Формула: total_price = cost_price * (1 + margin_percentage / 100)
        
        Args:
            cost_price: Себестоимость (сумма всех себестоимостей оборудования)
            margin_percentage: Процент маржи (если None, маржа не добавляется)
        
        Returns:
            Decimal: Итоговая цена
        """
        cost_price_decimal = Decimal(str(cost_price))
        
        if margin_percentage is None or margin_percentage == 0:
            return cost_price_decimal
        
        margin_decimal = Decimal(str(margin_percentage))
        total_price = cost_price_decimal * (Decimal('1') + margin_decimal / Decimal('100'))
        
        return total_price.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    
    @staticmethod
    def calculate_proposal_cost_price(proposal):
        """
        Рассчитать общую себестоимость КП на основе всех оборудования в EquipmentList.
        
        Args:
            proposal: CommercialProposal объект
        
        Returns:
            Decimal: Общая себестоимость
        """
        from .models import EquipmentListItem
        
        total_cost = Decimal('0')
        
        # Проходим по всем EquipmentList в КП
        for equipment_list in proposal.equipment_lists.all():
            # Проходим по всем EquipmentListItem
            for item in equipment_list.equipment_items_relation.all():
                equipment = item.equipment
                quantity = item.quantity
                
                # Рассчитываем себестоимость для этого оборудования
                try:
                    calculation_result = CostCalculationService.calculate_equipment_cost(
                        equipment=equipment,
                        exchange_rate_date=proposal.exchange_rate_date or None,
                        proposal=proposal,
                        target_currency=proposal.currency_ticket,
                        manual_overrides={'exchange_rate_value': proposal.exchange_rate} if proposal.exchange_rate else None
                    )
                    
                    # Себестоимость единицы оборудования в целевой валюте
                    unit_cost = Decimal(str(calculation_result['total_cost_target_currency']))
                    
                    # Учитываем количество
                    item_total_cost = unit_cost * quantity
                    total_cost += item_total_cost
                    
                except Exception as e:
                    # Если не удалось рассчитать, пропускаем это оборудование
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f'Failed to calculate cost for equipment {equipment.equipment_id}: {str(e)}')
                    continue
        
        return total_cost.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)


class DataAggregatorService:
    """
    Service for preparing data for the Proposal Constructor (Frontend & PDF).
    Aggregates data from various models and calculates final prices using the 'dissolution' method.
    Enforces security by sanitizing internal costs and margins.
    """

    def __init__(self, proposal):
        self.proposal = proposal
        self.currency = proposal.currency_ticket
        # Transform internal_exchange_rates list to dict for faster lookup
        # Expected format: [{'currency_from': 'USD', 'rate_value': 450.0, ...}, ...]
        # Mapping: currency_code -> rate to KZT
        self.rates_map = {}
        if self.proposal.internal_exchange_rates:
            for rate in self.proposal.internal_exchange_rates:
                if isinstance(rate, dict):
                    # Try different keys that might be used
                    curr = rate.get('currency_from') or rate.get('currency')
                    val = rate.get('rate_value') or rate.get('value') or rate.get('internal_rate')
                    if curr and val is not None:
                        self.rates_map[curr] = Decimal(str(val))
        
        # Ensure proposal exchange rate is also available
        if self.proposal.exchange_rate:
             # This is the rate of proposal.currency_ticket to KZT
             self.rates_map[self.proposal.currency_ticket] = Decimal(str(self.proposal.exchange_rate))
        
        # Add pivot currency
        self.rates_map['KZT'] = Decimal('1.0')

        # Fetch adjustments and calculate net percentage
        self.adjustments = self.proposal.adjustments.all()
        self.net_adjustment_percentage = Decimal('0')
        for adj in self.adjustments:
            val = Decimal(str(adj.value_percentage))
            if adj.adjustment_type == 'markup':
                self.net_adjustment_percentage += val
            else:
                self.net_adjustment_percentage -= val

    def get_full_data_package(self):
        """
        Main entry point. Returns a dictionary with all necessary data.
        """
        # Calculate equipment list first as it drives the core pricing
        equipment_list_data = self._calculate_and_build_equipment_list()
        
        # Now calculate adjustment values
        adjustment_value_kzt = self.total_adjusted_sale_price_sum - self.total_raw_sale_price_sum
        self.adjustment_value_target = self._convert_currency(adjustment_value_kzt, 'KZT', self.currency)
        self.adjustment_value_kzt = adjustment_value_kzt
        
        # Calculate total margin from all equipment items
        total_margin_kzt = Decimal('0')
        total_base_cost_kzt = Decimal('0')
        for item in equipment_list_data:
            if 'margin_kzt_total' in item:
                total_margin_kzt += Decimal(str(item['margin_kzt_total']))
            if 'base_cost_kzt' in item and 'quantity' in item:
                total_base_cost_kzt += Decimal(str(item['base_cost_kzt'])) * Decimal(str(item['quantity']))
        
        # Calculate total margin percentage
        total_margin_percentage = Decimal('0')
        if total_base_cost_kzt > 0:
            total_margin_percentage = (total_margin_kzt / total_base_cost_kzt) * Decimal('100')

        # Clamp and quantize margin values to fit model fields (max_digits=5, decimal_places=2)
        # Avoid InvalidOperation when serializing
        try:
            total_margin_kzt = total_margin_kzt.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        except Exception:
            total_margin_kzt = Decimal('0')

        try:
            # Max value that fits in DecimalField(max_digits=5, decimal_places=2) = 999.99
            max_margin_percentage = Decimal('999.99')
            if total_margin_percentage > max_margin_percentage:
                total_margin_percentage = max_margin_percentage
            total_margin_percentage = total_margin_percentage.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        except Exception:
            total_margin_percentage = Decimal('0')
        
        # Update proposal margin values
        self.proposal.margin_value = total_margin_kzt
        self.proposal.margin_percentage = total_margin_percentage
        self.proposal.save(update_fields=['margin_value', 'margin_percentage'])

        return {
            "proposal": self._get_proposal_metadata(),
            "client": self._get_client_data(),
            "user": self._get_user_data(),
            "equipment_list": equipment_list_data,
            "equipment_details": self._get_equipment_details(),
            "equipment_specifications": self._get_equipment_specifications(),
            "tech_processes": self._get_tech_processes(),
            "additional_services": self._get_additional_services(),
            "company_logo_url": self._get_company_logo_url(),
            "adjustments": self._get_adjustment_data(),
        }

    def _get_company_logo_url(self):
        sys_settings = SystemSettings.get_settings()
        if sys_settings.company_logo:
             return sys_settings.company_logo.url
        return "/static/assets/prosnab_logo.png"

    def _get_proposal_metadata(self):
        # Convert dates to strings for JSON serialization
        from datetime import datetime, date
        
        proposal_date_str = None
        if self.proposal.proposal_date:
            if isinstance(self.proposal.proposal_date, str):
                # Try to parse and reformat if it's in ISO format or other formats
                try:
                    # Try ISO format first (e.g., "2026-01-19" or "2026-01-19T20:12:50Z")
                    date_str = self.proposal.proposal_date.replace('Z', '+00:00')
                    if 'T' in date_str:
                        dt = datetime.fromisoformat(date_str)
                    else:
                        # Just date part (e.g., "2026-01-19")
                        dt = datetime.strptime(date_str.split('T')[0].split('+')[0], '%Y-%m-%d')
                    proposal_date_str = dt.strftime('%d.%m.%Y')
                except (ValueError, AttributeError):
                    # If parsing fails, check if it's already in d.m.Y format
                    try:
                        # Try to parse as d.m.Y to validate
                        datetime.strptime(self.proposal.proposal_date, '%d.%m.%Y')
                        proposal_date_str = self.proposal.proposal_date
                    except (ValueError, AttributeError):
                        # If all parsing fails, use as-is (might be in another format)
                        proposal_date_str = self.proposal.proposal_date
            elif isinstance(self.proposal.proposal_date, (date, datetime)):
                # Format as d.m.Y for template compatibility
                if isinstance(self.proposal.proposal_date, datetime):
                    proposal_date_str = self.proposal.proposal_date.strftime('%d.%m.%Y')
                else:
                    proposal_date_str = self.proposal.proposal_date.strftime('%d.%m.%Y')
            else:
                proposal_date_str = str(self.proposal.proposal_date) if self.proposal.proposal_date else None
        
        valid_until_str = None
        if self.proposal.valid_until:
            if isinstance(self.proposal.valid_until, str):
                # Try to parse and reformat if it's in ISO format
                try:
                    dt = datetime.fromisoformat(self.proposal.valid_until.replace('Z', '+00:00'))
                    valid_until_str = dt.strftime('%d.%m.%Y')
                except (ValueError, AttributeError):
                    valid_until_str = self.proposal.valid_until
            elif isinstance(self.proposal.valid_until, (date, datetime)):
                # Format as d.m.Y for template compatibility
                valid_until_str = self.proposal.valid_until.strftime('%d.%m.%Y')
            else:
                valid_until_str = str(self.proposal.valid_until)
        
        return {
            "id": self.proposal.proposal_id,
            "number": self.proposal.outcoming_number,
            "date": proposal_date_str,
            "valid_until": valid_until_str,
            "header": self.proposal.proposal_name,
            "currency": self.currency,
            "delivery_time": self.proposal.delivery_time,
            "warranty": self.proposal.warranty,
            "margin_value": float(self.proposal.margin_value) if self.proposal.margin_value else None,
            "total_price": float(self.proposal.total_price) if self.proposal.total_price else 0,
        }

    def _get_client_data(self):
        client = self.proposal.client
        if not client:
            return {}
        return {
            "name": client.client_name,
            "company": client.client_company_name,
            "address": client.client_address,
            "phone": client.client_phone,
            "email": client.client_email,
        }

    def _get_user_data(self):
        user = self.proposal.user
        if not user:
            return {}
        return {
            "name": user.user_name,
            "email": user.user_email,
            "phone": user.user_phone,
        }

    def _get_adjustment_data(self):
        """Return information about markups/discounts."""
        return {
            "net_percentage": float(self.net_adjustment_percentage),
            "value_kzt": float(self.adjustment_value_kzt),
            "value_target": float(self.adjustment_value_target),
            "type": "markup" if self.net_adjustment_percentage >= 0 else "discount",
            "items": [
                {
                    "type": adj.adjustment_type,
                    "value": float(adj.value_percentage),
                    "comments": adj.comments,
                    "author": adj.created_by.user_name if adj.created_by else "System"
                } for adj in self.adjustments
            ]
        }

    def _calculate_and_build_equipment_list(self):
        """
        New pricing method: Fixed sale price per unit with margin calculation.
        Step 1: Get sale_price_kzt from Equipment (fixed selling price).
        Step 2: Calculate Base Cost for each item (Purchase + Row Expenses).
        Step 3: Calculate Global Overhead (Total Proposal Price - Sum(Sale Prices)).
        Step 4: Distribute Global Overhead proportionally to Base Costs.
        Step 5: Calculate margin for each line: margin_kzt = sale_price_kzt - (base_cost + allocated_overhead_per_unit).
        Step 6: Build final list with sale_price_kzt as price_per_unit and calculated margins.
        """
        items_meta = []
        total_base_cost_sum = Decimal('0')
        total_sale_price_sum = Decimal('0')
        total_raw_sale_price_sum = Decimal('0')
        
        # Step 1: Collect all equipment items and calculate base costs
        for eq_list in self.proposal.equipment_lists.all():
            for item in eq_list.equipment_items_relation.select_related('equipment').prefetch_related('equipment__photos').all().order_by('order', 'created_at'):
                equipment = item.equipment
                quantity = Decimal(item.quantity)
                
                # Check if we have saved calculated_data
                saved_calculated_data = item.calculated_data or {}
                has_saved_data = (
                    saved_calculated_data.get('purchase_price_kzt') is not None or
                    saved_calculated_data.get('base_cost_kzt') is not None or
                    saved_calculated_data.get('margin_kzt') is not None
                )
                
                # Get sale_price_kzt (fixed selling price)
                sale_price_kzt = None
                if equipment.sale_price_kzt:
                    sale_price_kzt = Decimal(str(equipment.sale_price_kzt))
                # Also check saved price_per_unit as fallback
                if sale_price_kzt is None and item.price_per_unit:
                    sale_price_kzt = Decimal(str(item.price_per_unit))
                
                # If no sale_price_kzt, fallback to old method (for backward compatibility)
                if sale_price_kzt is None:
                    # Use old calculation method
                    purchase_price_obj = equipment.purchase_prices.filter(is_active=True).order_by('-created_at').first()
                    base_unit_cost = Decimal('0')
                    
                    if purchase_price_obj and purchase_price_obj.price:
                        price_val = Decimal(str(purchase_price_obj.price))
                        price_curr = purchase_price_obj.currency
                        base_unit_cost = self._convert_currency(price_val, price_curr, 'KZT')
                    elif equipment.equipment_manufacture_price:
                        price_val = Decimal(str(equipment.equipment_manufacture_price))
                        price_curr = equipment.equipment_price_currency_type or 'KZT'
                        base_unit_cost = self._convert_currency(price_val, price_curr, 'KZT')
                    
                    # Add row expenses
                    row_expenses_sum = Decimal('0')
                    if item.row_expenses and isinstance(item.row_expenses, list):
                        for exp in item.row_expenses:
                            val = exp.get('value')
                            curr = exp.get('currency', 'KZT')
                            if val:
                                row_expenses_sum += self._convert_currency(Decimal(str(val)), curr, 'KZT')
                    
                    if quantity > 0:
                        base_unit_cost += (row_expenses_sum / quantity)
                    
                    # Use calculated price as sale price (old behavior)
                    sale_price_kzt = base_unit_cost
                
                # Capture raw price for adjustment calculation
                raw_sale_price_kzt = sale_price_kzt
                
                # Round RAW line total to integer
                line_raw_total = (raw_sale_price_kzt * quantity).quantize(Decimal('1'), rounding=ROUND_HALF_UP)
                total_raw_sale_price_sum += line_raw_total

                # Apply Proposal Adjustments (Markups/Discounts) to the sale price
                if self.net_adjustment_percentage != 0:
                    sale_price_kzt = raw_sale_price_kzt * (Decimal('1') + self.net_adjustment_percentage / Decimal('100'))

                # Use saved calculated_data if available, otherwise calculate
                if has_saved_data:
                    # Use saved values from calculated_data (safely convert to Decimal)
                    try:
                        purchase_price_val = saved_calculated_data.get('purchase_price_kzt', 0)
                        if purchase_price_val is None:
                            purchase_price_kzt_for_save = Decimal('0')
                        else:
                            purchase_price_kzt_for_save = Decimal(str(purchase_price_val))
                    except (ValueError, TypeError, InvalidOperation):
                        purchase_price_kzt_for_save = Decimal('0')
                    
                    try:
                        base_cost_val = saved_calculated_data.get('base_cost_kzt', 0)
                        if base_cost_val is None:
                            base_unit_cost_kzt = Decimal('0')
                        else:
                            base_unit_cost_kzt = Decimal(str(base_cost_val))
                    except (ValueError, TypeError, InvalidOperation):
                        base_unit_cost_kzt = Decimal('0')
                else:
                    # Calculate base cost (purchase price + row expenses)
                    purchase_price_obj = equipment.purchase_prices.filter(is_active=True).order_by('-created_at').first()
                    base_unit_cost_kzt = Decimal('0')
                    
                    if purchase_price_obj and purchase_price_obj.price:
                        price_val = Decimal(str(purchase_price_obj.price))
                        price_curr = purchase_price_obj.currency
                        base_unit_cost_kzt = self._convert_currency(price_val, price_curr, 'KZT')
                    elif equipment.equipment_manufacture_price:
                        price_val = Decimal(str(equipment.equipment_manufacture_price))
                        price_curr = equipment.equipment_price_currency_type or 'KZT'
                        base_unit_cost_kzt = self._convert_currency(price_val, price_curr, 'KZT')
                    
                    # Add row expenses per unit
                    row_expenses_sum_kzt = Decimal('0')
                    if item.row_expenses and isinstance(item.row_expenses, list):
                        for exp in item.row_expenses:
                            val = exp.get('value')
                            curr = exp.get('currency', 'KZT')
                            if val:
                                row_expenses_sum_kzt += self._convert_currency(Decimal(str(val)), curr, 'KZT')
                    
                    if quantity > 0:
                        base_unit_cost_kzt += (row_expenses_sum_kzt / quantity)
                    
                    # Calculate purchase_price_kzt (before row expenses) for saving
                    purchase_price_kzt_for_save = Decimal('0')
                    if purchase_price_obj and purchase_price_obj.price:
                        price_val = Decimal(str(purchase_price_obj.price))
                        price_curr = purchase_price_obj.currency
                        purchase_price_kzt_for_save = self._convert_currency(price_val, price_curr, 'KZT')
                    elif equipment.equipment_manufacture_price:
                        price_val = Decimal(str(equipment.equipment_manufacture_price))
                        price_curr = equipment.equipment_price_currency_type or 'KZT'
                        purchase_price_kzt_for_save = self._convert_currency(price_val, price_curr, 'KZT')
                
                # Calculate adjusted totals
                # The frontend rounds the LINE total: Math.round(unit * (1+adj) * qty)
                # Let's match that.
                line_sale_total = (sale_price_kzt * quantity).quantize(Decimal('1'), rounding=ROUND_HALF_UP)
                
                line_base_total = base_unit_cost_kzt * quantity
                
                total_base_cost_sum += line_base_total
                total_sale_price_sum += line_sale_total
                
                items_meta.append({
                    'item': item,
                    'equipment': equipment,
                    'quantity': quantity,
                    'base_unit_cost_kzt': base_unit_cost_kzt,
                    'line_base_total': line_base_total,
                    'sale_price_kzt': sale_price_kzt,
                    'line_sale_total': line_sale_total,
                    'purchase_price_kzt': purchase_price_kzt_for_save,  # Save purchase price in KZT
                    'saved_calculated_data': saved_calculated_data,  # Pass saved data for later use
                })
        
        # Step 2: Calculate Global Overhead from EquipmentList
        # Get global expenses (tax, delivery, additional_prices) from EquipmentList
        global_overhead = Decimal('0')
        for eq_list in self.proposal.equipment_lists.all():
            # Tax
            if eq_list.tax_price:
                global_overhead += Decimal(str(eq_list.tax_price))
            # Delivery
            if eq_list.delivery_price:
                global_overhead += Decimal(str(eq_list.delivery_price))
            # Additional prices (if any)
            for add_price in eq_list.additional_prices.all():
                if add_price.value_type == 'fixed':
                    global_overhead += Decimal(str(add_price.price_parameter_value))
                elif add_price.value_type == 'percentage':
                    # Calculate percentage of base cost sum
                    if total_base_cost_sum > 0:
                        percentage_value = (Decimal(str(add_price.price_parameter_value)) / Decimal('100')) * total_base_cost_sum
                        global_overhead += percentage_value
        
        # Calculate services sum
        services_sum = Decimal('0')
        if self.proposal.additional_services:
            for svc in self.proposal.additional_services:
                price = svc.get('price')
                if price:
                    # Round each service price to integer
                    services_sum += Decimal(str(price)).quantize(Decimal('1'), rounding=ROUND_HALF_UP)
        
        # Calculate total price: sum of rounded sale prices + rounded services
        # We always recalculate this to ensure consistency with the new rounding rules
        calculated_total_kzt = total_sale_price_sum + services_sum
        
        # Convert total to target currency and round
        target_total = self._convert_currency(calculated_total_kzt, 'KZT', self.currency).quantize(Decimal('1'), rounding=ROUND_HALF_UP)
        
        # Update proposal total_price if it's different (to sync legacy data)
        if not self.proposal.total_price or abs(Decimal(str(self.proposal.total_price)) - target_total) > 0.01:
            self.proposal.total_price = target_total
            self.proposal.save(update_fields=['total_price'])
        
        # Step 3: Distribute Overhead and Calculate Margins
        final_items = []
        total_quantity = sum(meta['quantity'] for meta in items_meta) if items_meta else Decimal('1')
        
        for meta in items_meta:
            equipment = meta['equipment']
            quantity = meta['quantity']
            base_unit_cost_kzt = meta['base_unit_cost_kzt']
            sale_price_kzt = meta['sale_price_kzt']
            purchase_price_kzt = meta.get('purchase_price_kzt', Decimal('0'))  # Get from meta (already in KZT)
            saved_calculated_data = meta.get('saved_calculated_data', {})
            
            # Check if we should use saved calculated_data
            has_saved_margin = (
                saved_calculated_data and isinstance(saved_calculated_data, dict) and
                (saved_calculated_data.get('margin_kzt') is not None or
                 saved_calculated_data.get('margin_percentage') is not None)
            )
            
            if has_saved_margin:
                # Use saved values from calculated_data (safely convert to Decimal)
                try:
                    allocated_overhead_val = saved_calculated_data.get('allocated_overhead_per_unit', 0)
                    if allocated_overhead_val is None:
                        allocated_overhead_per_unit = Decimal('0')
                    else:
                        allocated_overhead_per_unit = Decimal(str(allocated_overhead_val))
                except (ValueError, TypeError, InvalidOperation):
                    allocated_overhead_per_unit = Decimal('0')
                
                try:
                    margin_kzt_val = saved_calculated_data.get('margin_kzt', 0)
                    if margin_kzt_val is None:
                        margin_kzt_per_unit = Decimal('0')
                    else:
                        margin_kzt_per_unit = Decimal(str(margin_kzt_val))
                except (ValueError, TypeError, InvalidOperation):
                    margin_kzt_per_unit = Decimal('0')
                
                try:
                    margin_perc_val = saved_calculated_data.get('margin_percentage', 0)
                    if margin_perc_val is None:
                        margin_percentage = Decimal('0')
                    else:
                        margin_percentage = Decimal(str(margin_perc_val))
                except (ValueError, TypeError, InvalidOperation):
                    margin_percentage = Decimal('0')
            else:
                # Distribute overhead proportionally to base costs
                weight = Decimal('0')
                if total_base_cost_sum > 0:
                    weight = meta['line_base_total'] / total_base_cost_sum
                elif total_quantity > 0:
                    weight = quantity / total_quantity
                
                # Allocated overhead per unit
                allocated_overhead_total = global_overhead * weight
                allocated_overhead_per_unit = Decimal('0')
                if quantity > 0:
                    allocated_overhead_per_unit = allocated_overhead_total / quantity
                
                # Calculate margin: sale_price - (base_cost + overhead)
                # Margin KZT per unit
                margin_kzt_per_unit = sale_price_kzt - (base_unit_cost_kzt + allocated_overhead_per_unit)
                
                # Margin percentage
                margin_percentage = Decimal('0')
                if base_unit_cost_kzt > 0:
                    margin_percentage = (margin_kzt_per_unit / base_unit_cost_kzt) * Decimal('100')
            
            # Total margin for the line
            margin_kzt_total = margin_kzt_per_unit * quantity
            
            # ROUND sale_price_kzt to nearest whole number
            sale_price_kzt = sale_price_kzt.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
            
            # Convert sale_price to target currency and ROUND to nearest whole number
            sale_price_target = self._convert_currency(sale_price_kzt, 'KZT', self.currency)
            sale_price_target = sale_price_target.quantize(Decimal('1'), rounding=ROUND_HALF_UP)

            # Final prices (sale_price_kzt is fixed, used as price_per_unit)
            # Note: base_cost_kzt and margins are kept in KZT for internal analysis
            final_items.append({
                "equipment_id": equipment.equipment_id,
                "name": equipment.equipment_name,
                "description": equipment.equipment_short_description,
                "article": equipment.equipment_articule,
                "quantity": int(quantity),
                "unit": equipment.equipment_uom or 'шт',
                "price_per_unit": float(sale_price_target),  # Fixed sale price in target currency
                "total_price": float(sale_price_target * quantity),  # Fixed sale price * quantity in target currency
                "base_cost_kzt": float(base_unit_cost_kzt),  # Purchase price + row expenses (in KZT)
                "allocated_overhead_per_unit": float(allocated_overhead_per_unit),  # Distributed overhead (in KZT) or saved value
                "margin_kzt": float(margin_kzt_per_unit),  # Margin per unit (in KZT) or saved value
                "margin_percentage": float(margin_percentage),  # Margin percentage or saved value
                "margin_kzt_total": float(margin_kzt_total),  # Total margin for line (in KZT)
                "purchase_price_kzt": float(purchase_price_kzt),  # Purchase price in KZT (before row expenses)
                "images": self._process_images_for_equipment(equipment)
            })
        
        # Save totals for metadata
        self.total_raw_sale_price_sum = total_raw_sale_price_sum
        self.total_adjusted_sale_price_sum = total_sale_price_sum

        return final_items

    def _get_additional_services(self):
        """
        Returns additional services with prices converted to the target currency and rounded.
        """
        services = self.proposal.additional_services or []
        if self.currency == 'KZT' or not services:
            # Even if KZT, ensure they are rounded to integers for consistency
            rounded_services = []
            for svc in services:
                new_svc = svc.copy()
                price = Decimal(str(svc.get('price', 0)))
                new_svc['price'] = float(price.quantize(Decimal('1'), rounding=ROUND_HALF_UP))
                rounded_services.append(new_svc)
            return rounded_services
            
        converted_services = []
        for svc in services:
            new_svc = svc.copy()
            price_kzt = Decimal(str(svc.get('price', 0)))
            # Convert and round to integer
            price_target = self._convert_currency(price_kzt, 'KZT', self.currency).quantize(Decimal('1'), rounding=ROUND_HALF_UP)
            new_svc['price'] = float(price_target)
            converted_services.append(new_svc)
        return converted_services

    def _convert_currency(self, amount, from_curr, to_curr):
        if from_curr == to_curr:
            return amount
            
        # Convert to KZT first (Base)
        rate_from = self.rates_map.get(from_curr)
        rate_to = self.rates_map.get(to_curr)
        
        if not rate_from:
             # Fallback: try to find in DB if not in internal map (sanity check, though map should have it)
             # Ideally validation should happen before
             # For now assume 1.0 or log error? Better to assume 1.0 to avoid crash
             rate_from = Decimal('1.0')
        
        if not rate_to:
             rate_to = Decimal('1.0')
             
        # Amount (From) * Rate (From->KZT) = KZT
        amount_kzt = amount * rate_from
        
        # KZT / Rate (To->KZT) = Amount (To)
        amount_target = amount_kzt / rate_to
        return amount_target

    def _get_equipment_details(self):
        """Get equipment details for ALL equipment in proposal, even if empty."""
        data = {}
        for eq_list in self.proposal.equipment_lists.all():
            for item in eq_list.equipment_items_relation.select_related('equipment').prefetch_related('equipment__photos').all().order_by('order', 'created_at'):
                equipment = item.equipment
                details = equipment.details.all()
                # Always include equipment, even if no details (empty list)
                data[equipment.equipment_id] = [
                    {'name': d.detail_parameter_name, 'value': d.detail_parameter_value} 
                    for d in details
                ]
        return data

    def _get_equipment_specifications(self):
        """Get equipment specifications for ALL equipment in proposal, even if empty. Includes Video Link if present."""
        data = {}
        for eq_list in self.proposal.equipment_lists.all():
            for item in eq_list.equipment_items_relation.select_related('equipment').prefetch_related('equipment__photos').all().order_by('order', 'created_at'):
                equipment = item.equipment
                specs = equipment.specifications.all()
                # Always include equipment, even if no specs (empty list)
                specs_list = [
                    {'name': s.spec_parameter_name, 'value': s.spec_parameter_value} 
                    for s in specs
                ]
                
                # Add Video Link as a specification row if present
                if equipment.equipment_videolinks:
                    specs_list.append({
                        'name': f"Ссылка на видео: {equipment.equipment_videolinks}", 
                        'value': '',
                        'is_video_link': True
                    })
                    
                data[equipment.equipment_id] = specs_list
        return data

    def _get_tech_processes(self):
        """Get tech processes for ALL equipment in proposal, even if empty."""
        data = {}
        for eq_list in self.proposal.equipment_lists.all():
            for item in eq_list.equipment_items_relation.select_related('equipment').prefetch_related('equipment__photos').all().order_by('order', 'created_at'):
                equipment = item.equipment
                procs = equipment.tech_processes.all()
                # Always include equipment, even if no processes (empty list)
                data[equipment.equipment_id] = [
                    {'title': p.tech_name, 'value': p.tech_value, 'desc': p.tech_desc} 
                    for p in procs
                ]
        return data

    def _process_images_for_equipment(self, equipment):
        """
        Build images list for equipment: local EquipmentPhoto first (stable /media/photos/ URLs),
        then legacy equipment_imagelinks (external, direct link). Used in data_package and PDF/DOCX.
        """
        result = []
        for photo in equipment.photos.all().order_by('sort_order', 'pk'):
            url = photo.image.url if photo.image else ''
            if url:
                result.append({'name': photo.name or '', 'url': url})
        imagelinks = equipment.equipment_imagelinks
        if not imagelinks:
            return result
        if isinstance(imagelinks, list):
            for item in imagelinks:
                if isinstance(item, dict):
                    name = item.get('name', '')
                    url = item.get('url')
                    if url:
                        direct_url = LinkConverterService.get_direct_link(url)
                        result.append({'name': name, 'url': direct_url})
                elif isinstance(item, str):
                    direct_url = LinkConverterService.get_direct_link(item)
                    result.append({'name': '', 'url': direct_url})
        elif isinstance(imagelinks, str):
            links = [link.strip() for link in imagelinks.split(',') if link.strip()]
            for link in links:
                direct_url = LinkConverterService.get_direct_link(link)
                result.append({'name': '', 'url': direct_url})
        return result

    def _process_images(self, imagelinks):
        """
        Parses image links field. Returns list of dicts with 'name' and 'url'.
        Kept for backward compatibility when equipment is not available.
        """
        if not imagelinks:
            return []
        result = []
        if isinstance(imagelinks, list):
            for item in imagelinks:
                if isinstance(item, dict):
                    name = item.get('name', '')
                    url = item.get('url')
                    if url:
                        direct_url = LinkConverterService.get_direct_link(url)
                        result.append({'name': name, 'url': direct_url})
                elif isinstance(item, str):
                    direct_url = LinkConverterService.get_direct_link(item)
                    result.append({'name': '', 'url': direct_url})
        elif isinstance(imagelinks, str):
            links = [link.strip() for link in imagelinks.split(',') if link.strip()]
class ExportService:
    """
    Service for generating export files (PDF, DOCX) from ProposalTemplate.
    """
    def __init__(self, template):
        self.template = template
        self.proposal = template.proposal
        
        # Get fresh data for price updates, but preserve saved equipment_list from constructor
        aggregator = DataAggregatorService(self.proposal)
        fresh_data_pkg = aggregator.get_full_data_package()
        
        # Use saved data_package from proposal if available (preserves constructor order and changes)
        if self.proposal.data_package:
            self.data_pkg = self.proposal.data_package.copy()
            saved_equipment_list = self.data_pkg.get('equipment_list', [])
            fresh_equipment_list = fresh_data_pkg.get('equipment_list', [])
            fresh_map = {item['equipment_id']: item for item in fresh_equipment_list}
            
            if saved_equipment_list:
                merged_list = []
                for saved_item in saved_equipment_list:
                    eq_id = saved_item.get('equipment_id')
                    if eq_id in fresh_map:
                        merged_item = saved_item.copy()
                        merged_item['price_per_unit'] = fresh_map[eq_id]['price_per_unit']
                        merged_item['total_price'] = fresh_map[eq_id]['total_price']
                        merged_item['quantity'] = fresh_map[eq_id]['quantity']
                        if 'base_cost_kzt' in fresh_map[eq_id]:
                            merged_item['base_cost_kzt'] = fresh_map[eq_id]['base_cost_kzt']
                        if 'images' in fresh_map[eq_id]:
                            merged_item['images'] = fresh_map[eq_id]['images']
                        if 'unit' not in merged_item or not merged_item.get('unit'):
                            merged_item['unit'] = fresh_map[eq_id].get('unit', 'шт')
                        merged_list.append(merged_item)
                    else:
                        if 'unit' not in saved_item or not saved_item.get('unit'):
                            saved_item['unit'] = 'шт'
                        merged_list.append(saved_item)
                self.data_pkg['equipment_list'] = merged_list
            else:
                self.data_pkg['equipment_list'] = fresh_equipment_list
        else:
            self.data_pkg = fresh_data_pkg
        
        if not self.data_pkg.get('equipment_specifications'):
            self.data_pkg['equipment_specifications'] = aggregator._get_equipment_specifications()
        if not self.data_pkg.get('equipment_details'):
            self.data_pkg['equipment_details'] = aggregator._get_equipment_details()
        if not self.data_pkg.get('tech_processes'):
            self.data_pkg['tech_processes'] = aggregator._get_tech_processes()
        
        self.data_pkg['proposal'] = aggregator._get_proposal_metadata()
        
        for item in self.data_pkg.get('equipment_list', []):
            if 'unit' not in item or not item.get('unit'):
                item['unit'] = 'шт'
        
        self.header_data = template.header_data or {}
        self.layout_data = template.layout_data or []
        
        # Fallback for header info from SystemSettings if template has none
        if not self.header_data.get('kz_info') or not self.header_data.get('ru_info'):
            from .models import SystemSettings
            sys_settings = SystemSettings.get_settings()
            if not self.header_data.get('kz_info'):
                self.header_data['kz_info'] = getattr(sys_settings, 'header_kz_info', '')
            if not self.header_data.get('ru_info'):
                self.header_data['ru_info'] = getattr(sys_settings, 'header_ru_info', '')

    def generate_pdf_html(self):
        """Builds HTML for PDF generation using layout_data and data_pkg."""
        from django.template.loader import render_to_string
        
        logo_url_or_path = self.data_pkg.get('company_logo_url')
        final_logo_path = ""
        if logo_url_or_path:
            if logo_url_or_path.startswith('/media/'):
                media_rel = logo_url_or_path[len(settings.MEDIA_URL):] if logo_url_or_path.startswith(settings.MEDIA_URL) else logo_url_or_path[7:]
                final_logo_path = os.path.join(settings.MEDIA_ROOT, media_rel)
            elif logo_url_or_path.startswith('/static/'):
                static_rel = logo_url_or_path[len(settings.STATIC_URL):] if logo_url_or_path.startswith(settings.STATIC_URL) else logo_url_or_path[8:]
                final_logo_path = os.path.join(settings.BASE_DIR, 'static', static_rel)
            else:
                final_logo_path = logo_url_or_path

        if not final_logo_path or not os.path.exists(final_logo_path):
             final_logo_path = os.path.join(settings.BASE_DIR, 'static/assets/prosnab_logo.png')
        
        logo_url = 'file://' + final_logo_path

        blocks_html = []
        for block in self.layout_data:
            title = block.get('title', '')
            content = block.get('content', '')
            spacing = block.get('spacing', 15)
            col_widths = block.get('columnWidths', {})
            had_placeholder = False
            
            if '{total_price_table}' in content: title = ''
            
            if '{equipment_list}' in content:
                content = content.replace('{equipment_list}', self._get_equipment_table_html(col_widths))
                had_placeholder = True
            if '{total_price_table}' in content:
                content = content.replace('{total_price_table}', self._get_total_price_html())
                had_placeholder = True
            if '{additional_services_table}' in content:
                content = content.replace('{additional_services_table}', self._get_additional_services_html(col_widths))
                had_placeholder = True
            if '{equipment_specs}' in content or '{equipment_specification}' in content:
                placeholder = '{equipment_specs}' if '{equipment_specs}' in content else '{equipment_specification}'
                content = content.replace(placeholder, self._get_equipment_specs_html(col_widths))
                had_placeholder = True
            if '{equipment_details}' in content or '{equipment_detail}' in content:
                placeholder = '{equipment_details}' if '{equipment_details}' in content else '{equipment_detail}'
                content = content.replace(placeholder, self._get_equipment_details_html(col_widths))
                had_placeholder = True
            if '{equipment_tech_process}' in content:
                content = content.replace('{equipment_tech_process}', self._get_equipment_tech_process_html(col_widths))
                had_placeholder = True
            if '{equipment_photo_grid}' in content:
                content = content.replace('{equipment_photo_grid}', self._get_photo_grid_html())
                had_placeholder = True

            if not had_placeholder:
                normalized = str(content).replace('\r\n', '\n').replace('\r', '\n')
                content = normalized.replace('\n', '<br/>')

            blocks_html.append({'title': title, 'content': content, 'spacing': spacing})

        context = {
            'proposal': self.data_pkg.get('proposal', {}),
            'client': self.data_pkg.get('client', {}),
            'user': self.data_pkg.get('user', {}),
            'header_data': self.header_data,
            'logo_url': logo_url,
            'blocks': blocks_html,
        }
        return render_to_string('proposals/pdf_template_new.html', context)

    def _get_equipment_table_html(self, col_widths):
        items = self.data_pkg.get('equipment_list', [])
        currency = self.proposal.currency_ticket
        name_width = col_widths.get('name', 300)
        
        rows = []
        total_sum = Decimal('0')
        for i, item in enumerate(items, 1):
            price = Decimal(str(item.get('price_per_unit', 0)))
            qty = Decimal(str(item.get('quantity', 0)))
            row_total = price * qty
            total_sum += row_total
            
            rows.append(f"""
            <tr>
                <td style="text-align: center; border: 1px solid #333; padding: 4px; font-size: 10pt;">{i}</td>
                <td style="border: 1px solid #333; padding: 4px; font-size: 10pt;">
                    <div class="item-name"><b>{escape(str(item.get('name', '')))}</b></div>
                    <div style="font-size: 9pt; color: #333;">{escape(str(item.get('description', '')))}</div>
                    {f'<div style="font-size: 9pt; color: #666;">Арт: {escape(str(item.get("article")))}</div>' if item.get('article') else ''}
                </td>
                <td style="text-align: center; border: 1px solid #333; padding: 4px; font-size: 10pt;">{qty}</td>
                <td style="text-align: center; border: 1px solid #333; padding: 4px; font-size: 10pt;">{escape(str(item.get('unit', 'шт')))}</td>
                <td style="text-align: right; border: 1px solid #333; padding: 4px; font-size: 10pt;">{price:,.2f} {currency}</td>
                <td style="text-align: right; border: 1px solid #333; padding: 4px; font-size: 10pt;">{row_total:,.2f} {currency}</td>
            </tr>
            """)
        
        total_row = f"""<tr style="background-color: #f5f5f5; font-weight: bold;">
            <td colspan="5" style="text-align: right; padding: 6px; border: 1px solid #333;">Итого:</td>
            <td style="text-align: right; padding: 6px; border: 1px solid #333;">{total_sum:,.2f} {currency}</td>
        </tr>"""
        
        return f"""<table style="width: 100%; border-collapse: collapse; border: 1px solid #333;">
            <thead><tr style="background-color: #f2f2f2;">
                <th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">№</th>
                <th style="width: {name_width}px; border: 1px solid #333; padding: 4px; font-size: 10pt;">Наименование</th>
                <th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Кол-во</th>
                <th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Ед.</th>
                <th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Цена</th>
                <th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Сумма</th>
            </tr></thead><tbody>{''.join(rows)}{total_row}</tbody></table>"""

    def _get_total_price_html(self):
        p = self.data_pkg.get('proposal', {})
        total = float(p.get("total_price", 0))
        currency = p.get("currency", "")
        return f'<div style="text-align: right; font-weight: bold; font-size: 10pt; margin-top: 10px;">ИТОГО: {total:,.2f} {currency}</div>'

    def _get_additional_services_html(self, col_widths):
        services = self.data_pkg.get('additional_services', [])
        if not services: return ""
        currency = self.proposal.currency_ticket
        desc_width = col_widths.get('description', 400)
        rows = ""
        for s in services:
            title = s.get('description') or s.get('name', '') or ''
            rows += f'<tr><td style="width: {desc_width}px; border: 1px solid #333; padding: 4px; font-size: 10pt;">{title}</td><td style="text-align: right; border: 1px solid #333; padding: 4px; font-size: 10pt;">{float(s.get("price", 0)):,.2f} {currency}</td></tr>'
        return f'<table style="width: 100%; border-collapse: collapse;"><thead><tr style="background-color: #f2f2f2;"><th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Описание</th><th style="border: 1px solid #333; padding: 4px; font-size: 10pt;">Стоимость</th></tr></thead><tbody>{rows}</tbody></table>'

    def _resolve_image_path_for_pdf(self, url):
        if not url: return ''
        if url.startswith('/media/'):
            media_rel = url[len(settings.MEDIA_URL):] if url.startswith(settings.MEDIA_URL) else url[7:]
            local_path = os.path.join(settings.MEDIA_ROOT, media_rel)
            if os.path.exists(local_path): return 'file://' + local_path
        elif url.startswith('/static/'):
            static_rel = url[len(settings.STATIC_URL):] if url.startswith(settings.STATIC_URL) else url[8:]
            local_path = os.path.join(settings.BASE_DIR, 'static', static_rel)
            if os.path.exists(local_path): return 'file://' + local_path
        return url

    def _get_equipment_specs_html(self, col_widths):
        items = self.data_pkg.get('equipment_list', [])
        html = ""
        specs_dict = self.data_pkg.get('equipment_specifications', {})
        for item in items:
            eq_id = item.get('equipment_id')
            if not eq_id: continue
            specs = specs_dict.get(eq_id) or specs_dict.get(str(eq_id))
            if not specs: continue
            images = item.get('images', []) or []
            total_rows = len(specs)
            num_photos = len(images)
            base_rowspan = total_rows // num_photos if num_photos else 0
            remainder = total_rows % num_photos if num_photos else 0
            
            def image_cell_at_row(row_index):
                if num_photos == 0: return (total_rows, None) if row_index == 0 else (None, None)
                start = 0
                for i in range(num_photos):
                    r = base_rowspan + (1 if i < remainder else 0)
                    if row_index == start: return (r, images[i])
                    start += r
                return (None, None)
            
            html += f'<div style="margin-bottom: 15px;"><h3 style="font-size: 10pt; margin-bottom: 3px; border-bottom: 1px solid #eee;">{item["name"]}</h3><table style="width: 100%; border-collapse: collapse; table-layout: fixed;"><thead><tr style="background-color: #f9f9f9;"><th style="width: 35%; border: 1px solid #333; padding: 4px; font-size: 10pt;">Параметр</th><th style="width: 20%; border: 1px solid #333; padding: 4px; font-size: 10pt;">Значение</th><th style="width: 45%; border: 1px solid #333; padding: 4px; font-size: 10pt;">Изображение</th></tr></thead><tbody>'
            for row_index, s in enumerate(specs):
                spec_name = s.get('name') or s.get('spec_parameter_name', '')
                spec_value = s.get('value') or s.get('spec_parameter_value', '')
                is_video_link = s.get('is_video_link', False)
                rowspan, img = image_cell_at_row(row_index)
                img_td = ''
                if rowspan is not None:
                    if img and img.get('url'):
                        img_path = self._resolve_image_path_for_pdf(img.get('url', ''))
                        caption = (img.get('name') or '').strip()
                        caption_html = f'<span style="font-size: 8pt; color: #666; display: block; margin-top: 4px;">{caption}</span>' if caption else ''
                        img_td = f'<td style="width: 45%; border: 1px solid #333; padding: 4px; text-align: center; vertical-align: middle;" rowspan="{rowspan}"><img src="{img_path}" style="width: 95%; height: auto; object-fit: contain; display: block; margin: 0 auto;">{caption_html}</td>'
                    else:
                        img_td = f'<td style="width: 45%; border: 1px solid #333; padding: 4px; vertical-align: middle;" rowspan="{rowspan}"></td>'

                if is_video_link:
                    display_name = escape(str(spec_name))
                    if 'http' in spec_name:
                        parts = spec_name.split(': ', 1)
                        if len(parts) == 2:
                            label, url = parts
                            display_name = f'{escape(label)}: <a href="{url}" style="color: blue; text-decoration: underline;">{escape(url)}</a>'
                    html += f'<tr><td colspan="2" style="border: 1px solid #333; padding: 4px; font-size: 10pt;">{display_name}</td>{img_td}</tr>'
                else:
                    spec_value_html = f'<a href="{spec_value}" style="color: blue; text-decoration: underline;">{spec_value}</a>' if str(spec_value).startswith('http') else escape(str(spec_value))
                    html += f'<tr><td style="width: 35%; border: 1px solid #333; padding: 4px; font-size: 10pt;">{escape(str(spec_name))}</td><td style="width: 20%; border: 1px solid #333; padding: 4px; font-size: 10pt;">{spec_value_html}</td>{img_td}</tr>'
            html += '</tbody></table></div>'
        return html

    def _get_equipment_details_html(self, col_widths):
        param_width = col_widths.get('param', 200)
        items = self.data_pkg.get('equipment_list', [])
        html = ""
        details_dict = self.data_pkg.get('equipment_details', {})
        for item in items:
            eq_id = item.get('equipment_id')
            if not eq_id: continue
            details = details_dict.get(eq_id) or details_dict.get(str(eq_id))
            if not details: continue
            html += f'<div style="margin-bottom: 15px;"><h3 style="font-size: 10pt; margin-bottom: 3px; border-bottom: 1px solid #eee;">{item["name"]}</h3><table style="width: 100%; border-collapse: collapse;">'
            for d in details:
                html += f'<tr><td style="width: {param_width}px; border: 1px solid #333; padding: 4px; font-size: 10pt;">{escape(str(d.get("name", "")))}</td><td style="border: 1px solid #333; padding: 4px; font-size: 10pt;">{escape(str(d.get("value", "")))}</td></tr>'
            html += '</table></div>'
        return html

    def _get_equipment_tech_process_html(self, col_widths):
        items = self.data_pkg.get('equipment_list', [])
        html = ""
        processes_dict = self.data_pkg.get('tech_processes', {})
        for item in items:
            eq_id = item.get('equipment_id')
            if not eq_id: continue
            processes = processes_dict.get(eq_id) or processes_dict.get(str(eq_id))
            if not processes: continue
            html += f'<div style="margin-bottom: 20px;"><h3 style="font-size: 10pt; margin-bottom: 5px; border-bottom: 1px solid #eee;">{item["name"]}</h3>'
            for proc in processes:
                html += '<div style="margin-bottom: 5px; padding: 4px; border: 1px solid #ddd; background-color: #f9f9f9;">'
                title = proc.get('title') or proc.get('tech_name', '')
                value = proc.get('value') or proc.get('tech_value', '')
                desc = proc.get('desc') or proc.get('tech_desc', '')
                if title: html += f'<div style="font-weight: bold; margin-bottom: 2px; font-size: 10pt;">{title}</div>'
                if value: html += f'<div style="margin-bottom: 2px; font-size: 10pt;">{value}</div>'
                if desc: html += f'<div style="color: #666; font-size: 9pt; margin-top: 2px;">{desc}</div>'
                html += '</div>'
            html += '</div>'
        return html

    def _get_photo_grid_html(self):
        items = self.data_pkg.get('equipment_list', [])
        html = '<div class="photo-grids-container">'
        for item in items:
            images = item.get('images', [])
            if not images: continue
            html += f'<table style="width: 100%; border-collapse: collapse; margin-bottom: 20px; page-break-inside: avoid;"><thead><tr style="background-color: #f9f9f9;"><th colspan="2" style="border: 1px solid #333; padding: 8px; text-align: left; font-size: 10pt;">{item.get("name", "")}</th></tr></thead><tbody>'
            for i in range(0, len(images), 2):
                pair = images[i:i+2]
                html += '<tr>'
                for img in pair:
                    url = img.get('url', '')
                    img_path = self._resolve_image_path_for_pdf(url)
                    html += f'<td style="width: 50%; height: 70mm; border: 1px solid #333; text-align: center; vertical-align: middle; padding: 5px;"><img src="{img_path}" style="max-width: 100%; max-height: 60mm; object-fit: contain;"><div style="font-size: 9pt; margin-top: 5px; color: #666;">{img.get("name", "")}</div></td>'
                if len(pair) == 1: html += '<td style="width: 50%; border: 1px solid #333;"></td>'
                html += '</tr>'
            html += '</tbody></table>'
        return html + '</div>'

    def generate_docx(self):
        """Generates DOCX using docxtpl and base_proposal.docx template."""
        import io, re, os
        template_path = os.path.join(settings.BASE_DIR, 'proposals/templates/base_proposal.docx')
        if not os.path.exists(template_path): return self._generate_docx_manual_fallback()
        try:
            doc = DocxTemplate(template_path)
            proposal_data = self.data_pkg.get('proposal', {})
            def clean_header_html(html_text):
                if not html_text: return ""
                text = re.sub(r'<br\s*/?>', '\n', html_text)
                text = re.sub(r'</p>', '\n', text)
                return re.sub(r'<[^>]+>', '', text).strip()
                
            header_subdoc = doc.new_subdoc()
            header_table = header_subdoc.add_table(rows=1, cols=2)
            header_table.columns[0].width = Mm(65)
            header_table.columns[1].width = Mm(65)
            
            kz_info = clean_header_html(self.header_data.get('kz_info', ''))
            if kz_info:
                p_kz = header_table.cell(0, 0).paragraphs[0]
                p_kz.add_run(kz_info).font.size = Pt(9)
                
            ru_info = clean_header_html(self.header_data.get('ru_info', ''))
            if ru_info:
                p_ru = header_table.cell(0, 1).paragraphs[0]
                p_ru.add_run(ru_info).font.size = Pt(9)
            
            logo_image = None
            logo_url = self.data_pkg.get('company_logo_url')
            if logo_url:
                rel = logo_url[len(settings.MEDIA_URL):] if logo_url.startswith(settings.MEDIA_URL) else logo_url[7:]
                final_path = os.path.join(settings.MEDIA_ROOT, rel)
                if not os.path.exists(final_path): final_path = os.path.join(settings.BASE_DIR, 'static/assets/prosnab_logo.png')
                if os.path.exists(final_path):
                    logo_image = InlineImage(doc, final_path, width=Mm(40))
            
            context = {
                'number': proposal_data.get('number', ''),
                'date': proposal_data.get('date', ''),
                'header': proposal_data.get('header', ''),
                'kz_info_clean': header_subdoc,
                'ru_info_clean': '',
                'company_logo_image': logo_image,
                'blocks': []
            }
            for block in self.layout_data:
                content = block.get('content', '')
                subdoc = doc.new_subdoc()
                
                title_text = block.get('title', '')
                if title_text:
                    p_title = subdoc.add_paragraph(title_text)
                    p_title.style.font.bold = True
                    p_title.style.font.size = Pt(12)
                    p_title.style.font.color.rgb = RGBColor(0, 0, 0)
                    
                had_placeholder = False
                if '{equipment_list}' in content:
                    self._add_equipment_table_docx(subdoc)
                    had_placeholder = True
                if '{total_price_table}' in content:
                    self._add_total_price_docx(subdoc)
                    had_placeholder = True
                if '{additional_services_table}' in content:
                    self._add_additional_services_docx(subdoc)
                    had_placeholder = True
                if '{equipment_specs}' in content or '{equipment_specification}' in content:
                    self._add_equipment_specs_docx(subdoc)
                    had_placeholder = True
                if '{equipment_details}' in content or '{equipment_detail}' in content:
                    self._add_equipment_details_docx(subdoc)
                    had_placeholder = True
                if '{equipment_tech_process}' in content:
                    self._add_equipment_tech_process_docx(subdoc)
                    had_placeholder = True
                if '{equipment_photo_grid}' in content:
                    self._add_photo_grid_docx(subdoc)
                    had_placeholder = True
                
                if not had_placeholder:
                    parser = DocxHTMLParser(subdoc)
                    parser.feed(content)
                    
                    # Remove trailing empty paragraphs
                    for p in reversed(subdoc.paragraphs):
                        if not p.text.strip() and not any(r.text.strip() for r in p.runs if r.text):
                            p._element.getparent().remove(p._element)
                        else:
                            break
                
                context['blocks'].append({'title': '', 'content_processed': subdoc})
            
            doc.render(context)
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}", exc_info=True)
            return self._generate_docx_manual_fallback()

    def _set_cell_background(self, cell, fill):
        """Helper to add background shading to a docx table cell."""
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _add_equipment_table_docx(self, doc):
        items = self.data_pkg.get('equipment_list', [])
        currency = self.proposal.currency_ticket
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        widths = [Mm(10), Mm(60), Mm(15), Mm(15), Mm(30), Mm(30)]
        for idx, width in enumerate(widths):
            table.columns[idx].width = width
            
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(['№', 'Наименование', 'Кол-во', 'Ед.', 'Цена', 'Сумма']):
            hdr_cells[i].text = text
            hdr_cells[i].width = widths[i]
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            self._set_cell_background(hdr_cells[i], 'f2f2f2')
        
        total_sum = 0
        for i, item in enumerate(items, 1):
            row = table.add_row().cells
            for idx, width in enumerate(widths):
                row[idx].width = width
            row[0].text = str(i)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_name = row[1].paragraphs[0]
            p_name.add_run(item.get('name', '')).bold = True
            
            if item.get('description'):
                p_desc = row[1].add_paragraph(item.get('description'))
                p_desc.style.font.size = Pt(9)
                
            if item.get('article'):
                p_art = row[1].add_paragraph(f"Арт: {item.get('article')}")
                p_art.style.font.size = Pt(9)
                p_art.style.font.color.rgb = RGBColor(102, 102, 102)
                
            row[2].text = str(item.get('quantity', 0))
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row[3].text = item.get('unit', 'шт')
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            price = float(item.get('price_per_unit', 0))
            row[4].text = f"{price:,.2f} {currency}"
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            total = float(item.get('total_price', 0))
            total_sum += total
            row[5].text = f"{total:,.2f} {currency}"
            row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for c in row: 
                if c.paragraphs and c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(10)
        
        row = table.add_row().cells
        row[0].merge(row[4])
        row[0].text = "Итого:"
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[5].text = f"{total_sum:,.2f} {currency}"
        row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        self._set_cell_background(row[0], 'f5f5f5')
        self._set_cell_background(row[5], 'f5f5f5')
        
        row[0].paragraphs[0].runs[0].font.bold = True
        row[5].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[5].paragraphs[0].runs[0].font.size = Pt(10)

    def _add_total_price_docx(self, doc):
        p = self.data_pkg.get('proposal', {})
        total = float(p.get("total_price", 0))
        currency = p.get("currency", "")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"ИТОГО: {total:,.2f} {currency}")
        run.font.size = Pt(10)
        run.font.bold = True

    def _add_additional_services_docx(self, doc):
        services = self.data_pkg.get('additional_services', [])
        if not services: return
        currency = self.proposal.currency_ticket
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Описание'
        hdr[1].text = 'Стоимость'
        for c in hdr:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.size = Pt(10)
            self._set_cell_background(c, 'f2f2f2')
        for s in services:
            row = table.add_row().cells
            row[0].text = s.get('description') or s.get('name', '')
            row[1].text = f"{float(s.get('price', 0)):,.2f} {currency}"
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for c in row: c.paragraphs[0].runs[0].font.size = Pt(10)

    def _add_equipment_specs_docx(self, doc):
        items = self.data_pkg.get('equipment_list', [])
        specs_dict = self.data_pkg.get('equipment_specifications', {})
        for item in items:
            eq_id = item.get('equipment_id')
            specs = specs_dict.get(eq_id) or specs_dict.get(str(eq_id))
            if not specs: continue
            
            doc.add_heading(item.get('name', ''), level=3).style.font.size = Pt(10)
            images = item.get('images', [])
            has_images = len(images) > 0
            cols = 3 if has_images else 2
            
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Параметр'; hdr[1].text = 'Значение'
            if has_images: hdr[2].text = 'Изображение'
            
            for c in hdr: 
                c.paragraphs[0].runs[0].font.size = Pt(10)
                c.paragraphs[0].runs[0].font.bold = True
                self._set_cell_background(c, 'f9f9f9')
            
            total_rows = len(specs)
            num_photos = len(images)
            base_rowspan = total_rows // num_photos if num_photos else 0
            remainder = total_rows % num_photos if num_photos else 0
            
            def image_cell_at_row(row_index):
                if num_photos == 0: return (None, None)
                start = 0
                for i in range(num_photos):
                    r = base_rowspan + (1 if i < remainder else 0)
                    if row_index == start: return (r, images[i])
                    start += r
                return (None, None)

            table_rows = []
            for row_index, s in enumerate(specs):
                row = table.add_row().cells
                table_rows.append(row)
                
                is_video_link = s.get('is_video_link', False)
                if is_video_link:
                    row[0].merge(row[1])
                    row[0].text = s.get('name', '')
                    row[0].paragraphs[0].runs[0].font.size = Pt(10)
                else:
                    row[0].text = s.get('name') or s.get('spec_parameter_name', '')
                    row[1].text = str(s.get('value') or s.get('spec_parameter_value', ''))
                    row[0].paragraphs[0].runs[0].font.size = Pt(10)
                    if row[1].paragraphs and row[1].paragraphs[0].runs:
                        row[1].paragraphs[0].runs[0].font.size = Pt(10)
                    else:
                        row[1].add_paragraph().add_run().font.size = Pt(10)
            
            # Apply rowspans and insert images
            if has_images:
                for row_index, row in enumerate(table_rows):
                    rowspan, img = image_cell_at_row(row_index)
                    if rowspan is not None:
                        if rowspan > 1:
                            target_cell = table_rows[row_index + rowspan - 1][2]
                            row[2].merge(target_cell)
                        
                        if img and img.get('url'):
                            url = img.get('url', '')
                            rel = url[len(settings.MEDIA_URL):] if url.startswith(settings.MEDIA_URL) else url[7:]
                            path = os.path.join(settings.MEDIA_ROOT, rel)
                            if os.path.exists(path):
                                para = row[2].paragraphs[0]
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                para.add_run().add_picture(path, width=Inches(2.0))
                                caption = (img.get('name') or '').strip()
                                if caption:
                                    cap_para = row[2].add_paragraph(caption)
                                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cap_para.style.font.size = Pt(8)
                                    cap_para.style.font.color.rgb = RGBColor(102, 102, 102)

            doc.add_paragraph('')

    def _add_equipment_details_docx(self, doc, col_widths=None):
        items = self.data_pkg.get('equipment_list', [])
        details_dict = self.data_pkg.get('equipment_details', {})
        for item in items:
            eq_id = item.get('equipment_id')
            details = details_dict.get(eq_id) or details_dict.get(str(eq_id))
            if not details: continue
            doc.add_heading(item.get('name', ''), level=3).style.font.size = Pt(10)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Параметр'; hdr[1].text = 'Значение'
            for c in hdr: 
                c.paragraphs[0].runs[0].font.size = Pt(10)
                c.paragraphs[0].runs[0].font.bold = True
                self._set_cell_background(c, 'f9f9f9')
            for d in details:
                row = table.add_row().cells
                row[0].text = d.get('name') or d.get('detail_parameter_name', '')
                row[1].text = str(d.get('value') or d.get('detail_parameter_value', ''))
                for c in row: 
                    if c.paragraphs and c.paragraphs[0].runs:
                        c.paragraphs[0].runs[0].font.size = Pt(10)
            doc.add_paragraph('')

    def _add_equipment_tech_process_docx(self, doc, col_widths=None):
        items = self.data_pkg.get('equipment_list', [])
        processes_dict = self.data_pkg.get('tech_processes', {})
        for item in items:
            eq_id = item.get('equipment_id')
            processes = processes_dict.get(eq_id) or processes_dict.get(str(eq_id))
            if not processes: continue
            doc.add_heading(item.get('name', ''), level=3).style.font.size = Pt(10)
            for proc in processes:
                title = proc.get('title') or proc.get('tech_name', '')
                value = proc.get('value') or proc.get('tech_value', '')
                desc = proc.get('desc') or proc.get('tech_desc', '')
                if title: 
                    p = doc.add_paragraph(title)
                    p.style.font.size = Pt(10)
                    p.style.font.bold = True
                if value: doc.add_paragraph(value).style.font.size = Pt(10)
                if desc: 
                    p = doc.add_paragraph(desc)
                    p.style.font.size = Pt(9)
                    p.style.font.italic = True
                    p.style.font.color.rgb = RGBColor(102, 102, 102)
            doc.add_paragraph('')

    def _add_photo_grid_docx(self, doc):
        items = self.data_pkg.get('equipment_list', [])
        for item in items:
            images = item.get('images', [])
            if not images: continue
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].merge(hdr[1])
            hdr[0].text = item.get('name', '')
            hdr[0].paragraphs[0].runs[0].font.size = Pt(10)
            hdr[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(hdr[0], 'f9f9f9')
            
            for i in range(0, len(images), 2):
                row = table.add_row().cells
                for j, img in enumerate(images[i:i+2]):
                    cell = row[j]
                    url = img.get('url', '')
                    rel = url[len(settings.MEDIA_URL):] if url.startswith(settings.MEDIA_URL) else url[7:]
                    path = os.path.join(settings.MEDIA_ROOT, rel)
                    if os.path.exists(path):
                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.add_run().add_picture(path, width=Inches(2.8))
                        if img.get('name'):
                            cap = cell.add_paragraph(img.get('name'))
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.style.font.size = Pt(9)
                            cap.style.font.color.rgb = RGBColor(102, 102, 102)
            doc.add_paragraph('')

    def _generate_docx_manual_fallback(self):
        import io
        doc = Document()
        doc.add_heading('Коммерческое предложение', 0)
        self._add_equipment_table_docx(doc)
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream
